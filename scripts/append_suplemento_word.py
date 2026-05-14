"""Append implementation supplement to Word TCC (keeps existing body)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt


def add_para(doc, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(12)


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    src = root / "WordFinalGlicNutri.docx"
    out = root / "WordFinalGlicNutri-ATUALIZADO.docx"
    if not src.is_file():
        raise SystemExit(f"Missing {src}")

    doc = Document(str(src))
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading(
        "Suplemento: implementação técnica e evidências no repositório (2026)",
        level=1,
    )
    add_para(
        doc,
        "Este suplemento foi acrescentado para alinhar o texto do TCC ao estado atual do "
        "projeto GlicNutri no repositório de código. O corpo original (modelagem, diagramas "
        "e protótipos) permanece válido como base da especificação; abaixo consolidam-se a "
        "implementação, a pilha tecnológica, o módulo de aprendizado de máquina e o "
        "encaixe com requisitos académicos (Bento/Thayse) e LGPD.",
    )

    doc.add_heading("1. Implementação e arquitetura", level=2)
    add_para(
        doc,
        "O sistema encontra-se implementado como aplicação móvel/web com Expo (React Native), "
        "com autenticação e persistência no Supabase (PostgreSQL, Row Level Security, "
        "Storage para artefactos de auditoria). O código organiza-se por telas "
        "(paciente, nutricionista, administrador), serviços e componentes reutilizáveis.",
    )

    doc.add_heading("2. Tecnologias (síntese para o capítulo de resultados)", level=2)
    for line in (
        "Cliente: Expo / React Native (JavaScript).",
        "Backend de dados: Supabase (Auth, Postgres, políticas de acesso, Storage).",
        "Machine Learning: Python, scikit-learn, notebooks Jupyter; API REST com FastAPI "
        "(endpoints /health e /predict).",
        "Integração local: a app chama a API de ML em ambiente de demonstração (por exemplo "
        "uvicorn em localhost), com CORS habilitado para testes na versão web.",
    ):
        doc.add_paragraph("• " + line)

    doc.add_heading("3. Machine Learning (disciplina Thayse)", level=2)
    add_para(
        doc,
        "O dataset oficial agrega uma linha por paciente e por dia (paciente-dia), exportada "
        "do Supabase por script reprodutível. O ficheiro manifest (export_manifest.json) "
        "regista contagem de linhas, colunas e hash do CSV. O notebook "
        "glicnutri_ml_pipeline.ipynb cobre EDA, quatro tarefas (classificação, regressão, "
        "clusterização, similaridade), métricas e persistência de artefactos joblib. A API "
        "expõe POST /predict alinhado às features e metadados de treino; no cliente existe "
        "fluxo de teste na área do paciente (Previsão com IA).",
    )

    doc.add_heading("4. Requisitos Bento (rastreabilidade)", level=2)
    add_para(
        doc,
        "A pasta entregas/bento no repositório contém checklist dos 13 requisitos, roteiros "
        "de CRUD e validações, evidências de base de dados e texto para usabilidade e "
        "relatórios. A auditoria e logs administrativos encontram-se implementados "
        "(serviço de auditoria, painel admin, armazenamento de logs). Evidências visuais da "
        "Semana 2: entregas/bento/semana-2-auditoria/ (incluindo capturas em prints/).",
    )

    doc.add_heading("5. LGPD e segurança (reforço)", level=2)
    add_para(
        doc,
        "O sistema trata dados sensíveis de saúde: aplicam-se finalidade específica, "
        "minimização (registos necessários ao acompanhamento), controlo de acesso por perfil, "
        "auditoria de operações relevantes e boas práticas de não exposição de credenciais "
        "(variáveis de ambiente e ficheiros locais; o CSV exportado com dados reais não é "
        "versionado por defeito). Recomenda-se manter esta secção alinhada ao texto legal "
        "curto exigido pela disciplina e à política do Supabase do grupo.",
    )

    doc.add_heading("6. Diagrama ER e base de dados", level=2)
    add_para(
        doc,
        "Para atualizar a Figura do DER com o esquema evolutivo do projeto, pode usar-se o "
        "diagrama vertical em entregas/diagrama-glicnutri-a4-vertical.pdf (ou .html/.png) e "
        "as migrations em supabase/migrations/ como fonte das tabelas atuais.",
    )

    doc.add_heading("7. Correções editoriais rápidas no corpo original (revisão manual)", level=2)
    for line in (
        "Capa: separar corretamente os nomes «Mateus Hilberath Costa» e «Gustavo Ribeiro» "
        "(evitar junção «COSTAGUSTAVO»).",
        "Ano de Curitiba/defesa: alinhar a 2026 se for o caso da banca.",
        "Secção 1.2.2: corrigir «EDesenvolver» para «Desenvolver».",
        "Secção 1: corrigir «INTRODUÇÃo» para «INTRODUÇÃO».",
        "Resumo e Abstract: atualizar o trecho que fala apenas em «especificação» para "
        "mencionar a implementação do protótipo funcional, ML e API (ver ficheiro "
        "entregas/WordFinalGlicNutri-ATUALIZACOES-PARA-COLAR.md).",
        "Conclusão: substituir a frase final que fala só em «trabalhos futuros» de "
        "implementação pela versão que reconhece o código entregue e mantém validação "
        "clínica como trabalho futuro.",
        "Secção 2.9.1: o sumário cita «Tecnologias utilizadas»; alinhar o título do corpo ou "
        "o sumário para a mesma redação.",
    ):
        doc.add_paragraph("• " + line)

    doc.save(str(out))
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
